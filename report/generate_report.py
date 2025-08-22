from __future__ import annotations
import argparse
import os
from docx import Document
from docx.shared import Pt
from app.config import AppConfig


def build_report(output_path: str):
	cfg = AppConfig()
	doc = Document()
	title = doc.add_heading("COVID-19 Analytics Platform - Project Report", 0)
	title.runs[0].font.size = Pt(18)
	doc.add_paragraph(f"Student: {cfg.student_name}")
	doc.add_paragraph("This report summarizes the implementation across Snowflake, API, EDA, visualization, and analytics features.")
	doc.add_heading("Tasks", level=1)
	doc.add_paragraph("1. Snowflake Marketplace dataset and resource monitor setup.")
	doc.add_paragraph("2. Data exploration and Python augmentation.")
	doc.add_paragraph("3. NoSQL schema for comments.")
	doc.add_paragraph("4. Flask API integrating Snowflake and MongoDB with caching.")
	doc.add_paragraph("5. Interactive Dash visualization.")
	doc.add_paragraph("6. Time series forecasting and clustering.")
	doc.add_paragraph("7. Performance optimizations in Snowflake.")
	doc.add_paragraph("8. API caching for frequently requested data.")
	doc.add_paragraph("9. Pattern recognition with MATCH_RECOGNIZE.")
	doc.add_paragraph("10. Share structures via Snowflake Data Sharing.")
	
	doc.add_heading("Insights", level=1)
	doc.add_paragraph("Include your analytical insights here: infection trends, mortality patterns, and demographic breakdowns.")

	doc.save(output_path)


if __name__ == "__main__":
	parser = argparse.ArgumentParser()
	parser.add_argument("--out", required=True, help="Output DOCX path")
	args = parser.parse_args()
	build_report(args.out)